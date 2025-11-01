"""
Report Export Service
Provides Word and PDF export functionality with proper table formatting.
"""

import io
import re
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
import markdown2
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend
import matplotlib.pyplot as plt
from sqlalchemy.orm import Session


def export_to_word(markdown_content: str, title: str) -> io.BytesIO:
    """
    Convert markdown report to Word document with proper formatting.

    Args:
        markdown_content: The markdown content to convert
        title: The report title

    Returns:
        BytesIO object containing the Word document
    """
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parse markdown line by line
    lines = markdown_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # H1 Headers
        if line.startswith('# '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H2 Headers
        elif line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H3 Headers
        elif line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Horizontal rules
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 80)

        # Tables
        elif line.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Back up one since we'll increment at the end

            # Parse table
            if len(table_lines) >= 2:  # Header + separator minimum
                # Parse header
                header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # Skip separator line
                # Parse data rows (skip separator)
                data_rows = []
                for row_line in table_lines[2:]:  # Skip header and separator
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    data_rows.append(cells)

                # Create Word table
                if data_rows:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
                    table.style = 'Light Grid Accent 1'

                    # Add header
                    header_cells = table.rows[0].cells
                    for idx, header_text in enumerate(header):
                        header_cells[idx].text = header_text
                        # Make header bold
                        for paragraph in header_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Add data rows
                    for row_idx, row_data in enumerate(data_rows, start=1):
                        cells = table.rows[row_idx].cells
                        for col_idx, cell_data in enumerate(row_data):
                            # Remove markdown formatting from cell content
                            clean_text = cell_data.replace('**', '')
                            cells[col_idx].text = clean_text

        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Remove markdown formatting
            text = text.replace('**', '')
            doc.add_paragraph(text, style='List Bullet')

        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            # Remove markdown formatting
            text = text.replace('**', '')
            doc.add_paragraph(text, style='List Number')

        # Regular paragraphs
        else:
            # Remove markdown formatting
            text = line.replace('**', '')
            if text:
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        i += 1

    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def export_to_pdf(markdown_content: str, title: str) -> io.BytesIO:
    """
    Convert markdown report to PDF document with proper formatting.

    Args:
        markdown_content: The markdown content to convert
        title: The report title

    Returns:
        BytesIO object containing the PDF document
    """
    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=30,
        alignment=TA_LEFT,
    )

    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#34495E'),
        spaceAfter=12,
        spaceBefore=12,
    )

    h2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#34495E'),
        spaceAfter=10,
        spaceBefore=10,
    )

    h3_style = ParagraphStyle(
        'CustomH3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor('#34495E'),
        spaceAfter=8,
        spaceBefore=8,
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )

    # Parse markdown
    lines = markdown_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # H1 Headers
        if line.startswith('# '):
            text = line[2:].strip()
            elements.append(Paragraph(text, h1_style))
            elements.append(Spacer(1, 12))

        # H2 Headers
        elif line.startswith('## '):
            text = line[3:].strip()
            elements.append(Paragraph(text, h2_style))
            elements.append(Spacer(1, 10))

        # H3 Headers
        elif line.startswith('### '):
            text = line[4:].strip()
            elements.append(Paragraph(text, h3_style))
            elements.append(Spacer(1, 8))

        # Horizontal rules
        elif line.startswith('---') or line.startswith('***'):
            elements.append(Spacer(1, 12))

        # Tables
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            if len(table_lines) >= 2:
                # Parse header
                header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # Parse data rows
                data_rows = [[cell.strip().replace('**', '') for cell in row.split('|')[1:-1]]
                            for row in table_lines[2:]]

                if data_rows:
                    # Create table data
                    table_data = [header] + data_rows

                    # Create table
                    t = Table(table_data, repeatRows=1)

                    # Style the table
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4A90E2')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
                    ]))

                    elements.append(t)
                    elements.append(Spacer(1, 12))

        # Bullet points or regular text
        else:
            text = line.replace('**', '<b>').replace('**', '</b>')
            text = text.replace('- ', '• ')
            if text:
                elements.append(Paragraph(text, body_style))
                elements.append(Spacer(1, 6))

        i += 1

    # Build PDF
    doc.build(elements)
    output.seek(0)

    return output


def export_to_word_with_charts(
    markdown_content: str,
    title: str,
    db: Session,
    user_id: int,
    brand_id: Optional[int] = None
) -> io.BytesIO:
    """
    Convert markdown report to Word document with embedded chart images.

    Args:
        markdown_content: The markdown content to convert
        title: The report title
        db: Database session for fetching analytics data
        user_id: User ID for fetching analytics data
        brand_id: Optional brand ID for filtering data

    Returns:
        BytesIO object containing the Word document with charts
    """
    from app import analytics
    from app import models

    # Fetch analytics data for placeholder replacement
    sentiment_data = analytics.get_sentiment_breakdown(db, user_id=user_id, brand_id=brand_id) or {}
    sov_data = analytics.get_share_of_voice(db, user_id=user_id, brand_id=brand_id)

    # Handle share_of_voice being either dict or list
    if isinstance(sov_data, list):
        # If it's a list, we can't use it for placeholders
        brand_sov = 0
    elif isinstance(sov_data, dict):
        brand_sov = sov_data.get('brand_sov', 0)
    else:
        brand_sov = 0

    # Get brand name
    brand_name = "Your Brand"  # Default
    if brand_id:
        brand = db.query(models.BrandInfo).filter(models.BrandInfo.id == brand_id).first()
        if brand:
            brand_name = brand.brand_name

    # Calculate positive sentiment rate (very_positive + positive)
    positive_sentiment_rate = sentiment_data.get('very_positive_pct', 0) + sentiment_data.get('positive_pct', 0)

    # Replace placeholders in markdown content
    markdown_content = markdown_content.replace('{brand_name}', brand_name)
    markdown_content = markdown_content.replace('{positive_sentiment_rate}', str(positive_sentiment_rate))
    markdown_content = markdown_content.replace('{descriptor_match_rate}', str(sentiment_data.get('descriptor_match_rate', 0)))
    markdown_content = markdown_content.replace('{share_of_voice[\'brand_sov\']}', str(brand_sov))

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Spacer

    # Generate charts as images
    chart_images = _generate_chart_images(db, user_id, brand_id)

    # Parse markdown content directly (no extra sections)
    lines = markdown_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # H1 Headers
        if line.startswith('# '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H2 Headers
        elif line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H3 Headers
        elif line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Images (markdown format: ![alt text](path))
        elif line.startswith('![') and '](' in line:
            # Extract image key from markdown
            # Format: ![Description](chart_paths['key']) or ![Description](report_charts/filename.png)
            try:
                # Try to find the chart key between single quotes
                if "chart_paths['" in line or 'chart_paths["' in line:
                    key_start = line.find("['") + 2 if "['" in line else line.find('["') + 2
                    key_end = line.find("']") if "']" in line else line.find('"]')
                    if key_start > 1 and key_end > key_start:
                        chart_key = line[key_start:key_end]
                        if chart_key in chart_images:
                            doc.add_picture(chart_images[chart_key], width=Inches(5.5))
                            doc.add_paragraph()  # Add spacing after image
                else:
                    # Try to extract chart name from filename (e.g., "sentiment" from "..._sentiment_...png")
                    # Look for common chart types in the filename
                    chart_types = ['dashboard', 'mention_rate', 'share_of_voice', 'sentiment',
                                   'platform_comparison', 'positioning', 'descriptor_performance']
                    for chart_type in chart_types:
                        if chart_type in line.lower().replace('_', ' ') or chart_type.replace('_', ' ') in line.lower():
                            # Map filename patterns to chart keys
                            chart_key_map = {
                                'dashboard': 'dashboard',
                                'mention rate': 'mention_rate',
                                'share of voice': 'share_of_voice',
                                'sentiment': 'sentiment',
                                'platform': 'platform_comparison',
                                'positioning': 'positioning',
                                'descriptor': 'descriptor_performance'
                            }
                            for pattern, key in chart_key_map.items():
                                if pattern in line.lower():
                                    if key in chart_images:
                                        doc.add_picture(chart_images[key], width=Inches(5.5))
                                        doc.add_paragraph()  # Add spacing after image
                                    break
                            break
            except Exception as e:
                print(f"Error embedding image: {e}")
                # Skip the image if there's an error

        # Horizontal rules
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 80)

        # Tables
        elif line.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            # Parse table
            if len(table_lines) >= 2:
                header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                data_rows = []
                for row_line in table_lines[2:]:
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    data_rows.append(cells)

                if data_rows:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
                    table.style = 'Light Grid Accent 1'

                    # Add header
                    header_cells = table.rows[0].cells
                    for idx, header_text in enumerate(header):
                        header_cells[idx].text = header_text
                        for paragraph in header_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Add data rows
                    for row_idx, row_data in enumerate(data_rows, start=1):
                        cells = table.rows[row_idx].cells
                        for col_idx, cell_data in enumerate(row_data):
                            clean_text = cell_data.replace('**', '')
                            cells[col_idx].text = clean_text

        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip().replace('**', '')
            doc.add_paragraph(text, style='List Bullet')

        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s+', '', line).replace('**', '')
            doc.add_paragraph(text, style='List Number')

        # Regular paragraphs
        else:
            text = line.replace('**', '')
            if text:
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        i += 1

    # Clean up chart images
    for img_buffer in chart_images.values():
        img_buffer.close()

    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def _generate_chart_images(db: Session, user_id: int, brand_id: Optional[int] = None) -> dict:
    """Generate chart images for Word document embedding."""
    from app import analytics

    chart_images = {}

    # TALES brand colors
    COLORS = {
        'primary': '#665775',
        'secondary': '#80a1d4',
        'accent': '#75c9c8',
        'positive': '#4caf50',
        'neutral': '#9e9e9e',
        'negative': '#f44336',
    }

    try:
        # 1. Sentiment Pie Chart
        sentiment_data = analytics.get_sentiment_breakdown(db, user_id=user_id, brand_id=brand_id)
        if sentiment_data:
            fig, ax = plt.subplots(figsize=(8, 6))

            # Build sentiment breakdown from flat structure
            sentiment_breakdown = {}
            if sentiment_data.get('very_positive', 0) > 0:
                sentiment_breakdown['Very Positive'] = sentiment_data['very_positive']
            if sentiment_data.get('positive', 0) > 0:
                sentiment_breakdown['Positive'] = sentiment_data['positive']
            if sentiment_data.get('neutral', 0) > 0:
                sentiment_breakdown['Neutral'] = sentiment_data['neutral']
            if sentiment_data.get('negative', 0) > 0:
                sentiment_breakdown['Negative'] = sentiment_data['negative']
            if sentiment_data.get('very_negative', 0) > 0:
                sentiment_breakdown['Very Negative'] = sentiment_data['very_negative']
            if sentiment_data.get('mixed', 0) > 0:
                sentiment_breakdown['Mixed'] = sentiment_data['mixed']

            if sentiment_breakdown:
                sentiments = list(sentiment_breakdown.keys())
                counts = list(sentiment_breakdown.values())

                # Color mapping for sentiments
                sentiment_colors = {
                    'Very Positive': '#58A13B',  # Extended green
                    'Positive': '#B2C9AB',       # Sage
                    'Neutral': '#9FA8DA',        # Periwinkle
                    'Negative': '#E04320',       # Burnt orange/red
                    'Very Negative': '#EA4A4A',  # Extended red
                    'Mixed': '#75C9C8'           # Teal
                }
                colors_list = [sentiment_colors.get(s, '#999999') for s in sentiments]

                ax.pie(counts, labels=sentiments, autopct='%1.1f%%', colors=colors_list, startangle=90)
                ax.set_title('Sentiment Distribution', fontsize=16, fontweight='bold', pad=20)

                # Save to BytesIO
                img_buffer = io.BytesIO()
                plt.tight_layout()
                plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                img_buffer.seek(0)
                chart_images['sentiment'] = img_buffer
                plt.close(fig)

        # 2. Positioning Bar Chart
        positioning_data = analytics.get_positioning_breakdown(db, user_id=user_id, brand_id=brand_id)
        if positioning_data:
            fig, ax = plt.subplots(figsize=(8, 6))

            # Build positioning breakdown from flat structure
            positioning_breakdown = {}
            if positioning_data.get('leader', 0) > 0:
                positioning_breakdown['Leader'] = positioning_data['leader']
            if positioning_data.get('top_3', 0) > 0:
                positioning_breakdown['Top 3'] = positioning_data['top_3']
            if positioning_data.get('featured', 0) > 0:
                positioning_breakdown['Featured'] = positioning_data['featured']
            if positioning_data.get('listed', 0) > 0:
                positioning_breakdown['Listed'] = positioning_data['listed']
            if positioning_data.get('not_mentioned', 0) > 0:
                positioning_breakdown['Not Mentioned'] = positioning_data['not_mentioned']

            if positioning_breakdown:
                positions = list(positioning_breakdown.keys())
                counts = list(positioning_breakdown.values())

                bars = ax.barh(positions, counts, color=COLORS['primary'])
                ax.set_xlabel('Number of Responses', fontsize=12)
                ax.set_title('Brand Positioning Breakdown', fontsize=16, fontweight='bold', pad=20)
                ax.invert_yaxis()

                # Add value labels
                for i, (bar, count) in enumerate(zip(bars, counts)):
                    ax.text(count, i, f' {count}', va='center', fontsize=10)

                # Save to BytesIO
                img_buffer = io.BytesIO()
                plt.tight_layout()
                plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                img_buffer.seek(0)
                chart_images['positioning'] = img_buffer
                plt.close(fig)

        # 3. Share of Voice Chart
        sov_data = analytics.get_share_of_voice(db, user_id=user_id, brand_id=brand_id)
        if sov_data and isinstance(sov_data, list) and len(sov_data) > 0:
            fig, ax = plt.subplots(figsize=(8, 6))

            # Sort by total_mentions (descending)
            sorted_data = sorted(sov_data, key=lambda x: x.get('total_mentions', 0), reverse=True)

            # Extract organization names and mention counts
            entities = [item.get('organization', 'Unknown') for item in sorted_data]
            counts = [item.get('total_mentions', 0) for item in sorted_data]

            # Only include if there are counts > 0
            if any(c > 0 for c in counts):
                # Highlight brand in different color
                colors_list = [COLORS['primary'] if item.get('is_brand', False) else COLORS['secondary'] for item in sorted_data]

                bars = ax.barh(entities, counts, color=colors_list)
                ax.set_xlabel('Mention Count', fontsize=12)
                ax.set_title('Share of Voice - Brand vs Competitors', fontsize=16, fontweight='bold', pad=20)
                ax.invert_yaxis()

                # Add value labels
                for i, (bar, count) in enumerate(zip(bars, counts)):
                    ax.text(count, i, f' {count}', va='center', fontsize=10)

                # Save to BytesIO
                img_buffer = io.BytesIO()
                plt.tight_layout()
                plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                img_buffer.seek(0)
                chart_images['share_of_voice'] = img_buffer
                plt.close(fig)

        # 4. Mention Rate Pie Chart
        # Use dashboard_metrics to get mention data
        dashboard_data = analytics.get_dashboard_metrics(db, user_id=user_id, brand_id=brand_id)
        if dashboard_data and dashboard_data.get('total_responses', 0) > 0:
            # Calculate mentions breakdown
            yes_count = dashboard_data.get('mention_count', 0)
            total = dashboard_data.get('total_responses', 0)
            no_count = total - yes_count
            labels = ['Mentioned', 'Not Mentioned']
            sizes = [yes_count, no_count]
            colors_list = ['#58A13B', '#E04320']  # success, danger

            if any(s > 0 for s in sizes):
                fig, ax = plt.subplots(figsize=(8, 6))
                wedges, texts, autotexts = ax.pie(
                    sizes,
                    explode=(0.05, 0),
                    labels=labels,
                    colors=colors_list,
                    autopct='%1.1f%%',
                    startangle=90
                )

                # Make percentage text white and bold
                for autotext in autotexts:
                    autotext.set_color('white')
                    autotext.set_fontsize(12)
                    autotext.set_weight('bold')

                ax.set_title('Brand Mention Rate in AI Responses', fontsize=16, fontweight='bold', pad=20)

                # Save to BytesIO
                img_buffer = io.BytesIO()
                plt.tight_layout()
                plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                img_buffer.seek(0)
                chart_images['mention_rate'] = img_buffer
                plt.close(fig)

        # 5. Platform Comparison Chart
        platform_data = analytics.get_platform_breakdown(db, user_id=user_id, brand_id=brand_id)
        if platform_data and isinstance(platform_data, dict):
            platforms = [p for p, m in platform_data.items() if m.get('total', 0) > 0]

            if platforms:
                mention_rates = [platform_data[p]['mention'].get('yes_pct', 0) for p in platforms]
                positive_sentiment = [
                    platform_data[p]['sentiment'].get('positive_pct', 0) +
                    platform_data[p]['sentiment'].get('very_positive_pct', 0)
                    for p in platforms
                ]
                leader_positioning = [
                    platform_data[p]['positioning'].get('leader_pct', 0) +
                    platform_data[p]['positioning'].get('top_3_pct', 0)
                    for p in platforms
                ]

                x = range(len(platforms))
                width = 0.25

                fig, ax = plt.subplots(figsize=(10, 6))

                bars1 = ax.bar([i - width for i in x], mention_rates, width, label='Mention Rate (%)',
                              color=COLORS['primary'])
                bars2 = ax.bar(x, positive_sentiment, width, label='Positive Sentiment (%)',
                              color='#58A13B')
                bars3 = ax.bar([i + width for i in x], leader_positioning, width, label='Leader/Top 3 (%)',
                              color=COLORS['secondary'])

                # Add value labels on bars
                for bars in [bars1, bars2, bars3]:
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            ax.text(bar.get_x() + bar.get_width()/2., height,
                                  f'{height:.1f}%',
                                  ha='center', va='bottom', fontsize=9, fontweight='bold')

                ax.set_ylabel('Percentage (%)', fontsize=11, fontweight='bold')
                ax.set_xlabel('Platform', fontsize=11, fontweight='bold')
                ax.set_title('Performance by Platform', fontsize=16, fontweight='bold', pad=20)
                ax.set_xticks(x)
                ax.set_xticklabels(platforms, fontsize=10)
                ax.legend(fontsize=9, loc='upper right')
                ax.grid(axis='y', alpha=0.3, linestyle='--')

                # Save to BytesIO
                img_buffer = io.BytesIO()
                plt.tight_layout()
                plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                img_buffer.seek(0)
                chart_images['platform_comparison'] = img_buffer
                plt.close(fig)

        # 6. Descriptor Performance Chart
        descriptor_data = analytics.get_descriptor_analysis(db, user_id=user_id, brand_id=brand_id)
        if descriptor_data and isinstance(descriptor_data, dict):
            sorted_descriptors = sorted(descriptor_data.items(), key=lambda x: x[1], reverse=True)[:10]

            if sorted_descriptors:
                descriptors, counts = zip(*sorted_descriptors)

                fig, ax = plt.subplots(figsize=(10, 6))
                bars = ax.barh(range(len(descriptors)), counts, color=COLORS['primary'])

                # Add value labels
                for bar, count in zip(bars, counts):
                    width = bar.get_width()
                    ax.text(width + 0.3, bar.get_y() + bar.get_height()/2.,
                          f'{int(count)}',
                          ha='left', va='center', fontsize=10, fontweight='bold')

                ax.set_yticks(range(len(descriptors)))
                ax.set_yticklabels(descriptors, fontsize=10)
                ax.set_xlabel('Number of Mentions', fontsize=11, fontweight='bold')
                ax.set_title('Top Descriptors Associated with Brand', fontsize=16, fontweight='bold', pad=20)
                ax.grid(axis='x', alpha=0.3, linestyle='--')
                ax.invert_yaxis()  # Highest at top

                # Save to BytesIO
                img_buffer = io.BytesIO()
                plt.tight_layout()
                plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                img_buffer.seek(0)
                chart_images['descriptor_performance'] = img_buffer
                plt.close(fig)

    except Exception as e:
        print(f"Error generating charts: {e}")
        # Return empty dict if chart generation fails
        pass

    return chart_images
